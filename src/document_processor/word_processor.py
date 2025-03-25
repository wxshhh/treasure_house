"""Word文档处理模块，负责解析Word文档并提取文本内容"""

import os
import docx
from typing import List, Dict, Any, Optional, Callable
from config import DOCUMENT_CONFIG


from .base_processor import BaseProcessor

class WordProcessor(BaseProcessor):
    """Word文档处理器，用于解析Word文档并提取文本内容"""

    def __init__(self, chunk_size: int = None, chunk_overlap: int = None):
        """初始化Word处理器

        Args:
            chunk_size: 文档分块大小，默认使用配置文件中的设置
            chunk_overlap: 分块重叠大小，默认使用配置文件中的设置
        """
        super().__init__(chunk_size, chunk_overlap)

    def set_progress_callback(self, callback: Callable[[str, float], None]):
        """设置进度回调函数
        
        Args:
            callback: 回调函数，接受阶段名称和进度值两个参数
        """
        self.progress_callback = callback

    def extract_text(self, file_path: str) -> str:
        """从Word文件中提取全部文本

        Args:
            file_path: Word文件路径

        Returns:
            提取的文本内容
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        try:
            # 打开Word文件
            doc = docx.Document(file_path)
            text = ""

            # 提取段落文本
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text.strip() + "\n"

            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text += " | ".join(row_text) + "\n"

            return text
        except Exception as e:
            raise Exception(f"Word文件解析失败: {str(e)}")

    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """提取Word文件的元数据

        Args:
            file_path: Word文件路径

        Returns:
            元数据字典
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        try:
            doc = docx.Document(file_path)
            core_properties = doc.core_properties
            
            metadata = {
                "title": core_properties.title or "",
                "author": core_properties.author or "",
                "subject": core_properties.subject or "",
                "keywords": core_properties.keywords or "",
                "created": str(core_properties.created) if core_properties.created else "",
                "modified": str(core_properties.modified) if core_properties.modified else "",
                "last_modified_by": core_properties.last_modified_by or "",
                "paragraph_count": len(doc.paragraphs),
                "file_size": os.path.getsize(file_path),
            }
            return metadata
        except Exception as e:
            raise Exception(f"Word元数据提取失败: {str(e)}")



    def process(self, file_path: str) -> Dict[str, Any]:
        """处理Word文档，提取文本并分块

        Args:
            file_path: Word文件路径

        Returns:
            包含文本块和元数据的字典
        """
        if self.progress_callback:
            self.progress_callback("开始处理", 0.0)

        text = self.extract_text(file_path)
        if self.progress_callback:
            self.progress_callback("文本提取完成", 0.3)

        metadata = self.extract_metadata(file_path)
        if self.progress_callback:
            self.progress_callback("元数据提取完成", 0.6)

        chunks = self.chunk_text(text)
        if self.progress_callback:
            self.progress_callback("分块完成", 1.0)

        # 保存处理结果到实例变量
        self.result = {
            "chunks": chunks,
            "metadata": metadata,
            "source": file_path,
            "source_type": "file",
            "total_chunks": len(chunks),
        }
        return self.result

    def get_result(self) -> Optional[Dict[str, Any]]:
        """获取处理结果"""
        return getattr(self, 'result', None)
